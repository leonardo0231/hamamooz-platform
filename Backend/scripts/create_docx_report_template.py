"""Create the versioned, allowlisted Word template used by report drafts.

This is intentionally a build-time asset generator, not a user-supplied
template mechanism.  The template exposes only the frozen report snapshot.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

TARGET = Path(__file__).resolve().parents[1] / "templates" / "reports" / "report_card.docx"


def add_line(document, text, *, bold=False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def main():
    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "DejaVu Sans"
    normal.font.size = Pt(10)

    heading = document.add_heading("Official student report", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Official student report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_line(document, "{% for report in reports %}")
    add_line(document, '{% if "student_identity" in blocks %}')
    add_line(document, "Student: {{ report.student.full_name }}", bold=True)
    add_line(document, "Student number: {{ report.student.student_number }}")
    add_line(
        document, "Academic year: {{ report.academic.year }} · Term: {{ report.academic.term }}"
    )
    add_line(document, "Grade: {{ report.academic.grade }} · Class: {{ report.academic.class }}")
    add_line(
        document,
        "Overall average: {{ report.summary.average }} · Class rank: {{ report.summary.class_rank }}",
    )
    add_line(document, "{% endif %}")

    add_line(document, '{% if "academic_summary" in blocks %}')
    document.add_heading("Academic summary", level=1)
    add_line(document, "{% if overrides.academic_summary %}")
    add_line(document, "{{ overrides.academic_summary }}")
    add_line(document, "{% else %}")
    add_line(document, "{% for subject in report.subjects %}")
    add_line(
        document,
        "{{ subject.title }} — Average: {{ subject.average }} — Status: {{ subject.passed }}",
    )
    add_line(document, "{% endfor %}")
    add_line(document, "{% endif %}")
    add_line(document, "{% endif %}")

    add_line(document, '{% if "attendance_summary" in blocks %}')
    document.add_heading("Attendance", level=1)
    add_line(
        document,
        "Finalized sessions: {{ report.product_context.attendance.finalized_session_count }}",
    )
    add_line(
        document,
        "Unexcused absences: {{ report.product_context.attendance.unexcused_absence_count }}",
    )
    add_line(document, "{% endif %}")

    add_line(document, '{% if "recommendations" in blocks %}')
    document.add_heading("Approved recommendations", level=1)
    add_line(document, "{% if overrides.recommendations %}")
    add_line(document, "{{ overrides.recommendations }}")
    add_line(document, "{% else %}")
    add_line(
        document, "{% for recommendation in report.product_context.approved_recommendations %}"
    )
    add_line(document, "{{ recommendation.approved_text }}")
    add_line(document, "{% endfor %}")
    add_line(document, "{% endif %}")
    add_line(document, "{% endif %}")
    add_line(document, "{% endfor %}")
    document.add_paragraph("Generated from a frozen, approved report snapshot.")
    document.save(TARGET)


if __name__ == "__main__":
    main()
