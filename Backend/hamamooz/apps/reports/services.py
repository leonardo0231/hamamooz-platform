import hashlib
import json
import math
import uuid
from collections import defaultdict
from copy import deepcopy
from datetime import timedelta
from decimal import Decimal
from html import escape
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from django.db.models import Max, Q
from django.template.loader import render_to_string
from django.utils import timezone

from hamamooz.apps.academics.calculations import (
    calculate_enrollment_term,
    get_policy,
    normalized_score,
    quantize,
    recalculate_class_term,
)
from hamamooz.apps.academics.models import (
    Assessment,
    CourseOffering,
    Score,
    SubjectResult,
    TermResult,
)
from hamamooz.apps.students.models import Enrollment

from .models import (
    ReportArchive,
    ReportDraft,
    ReportLayoutKey,
    ReportPeriodType,
)

ALLOWED_REPORT_BLOCKS = {
    "student_identity",
    "academic_summary",
    "attendance_summary",
    "evaluation_radar",
    "strengths",
    "weaknesses",
    "recommendations",
    "signatures",
}

# The layout is data, but not executable template source.  Keeping the CSS
# values here prevents a manager-provided presentation JSON object from
# influencing @page with arbitrary text.
ALLOWED_REPORT_PAGE_SIZES = {
    "a4_portrait": "A4 portrait",
    "a3_landscape": "A3 landscape",
}

REPORT_CARD_LAYOUTS = {
    "analytical_term_1": {
        "period_type": "term",
        "term_code": "first",
        "page_size": "a3_landscape",
        "template": "reports/analytical_term_1.html",
    },
    "analytical_term_2": {
        "period_type": "term",
        "term_code": "second",
        "page_size": "a3_landscape",
        "template": "reports/analytical_term_2.html",
    },
    "analytical_annual": {
        "period_type": "annual",
        "page_size": "a3_landscape",
        "template": "reports/analytical_annual.html",
    },
    "final_term_1": {
        "period_type": "term",
        "term_code": "first",
        "page_size": "a4_portrait",
        "template": "reports/final_term_1.html",
    },
    "final_term_2": {
        "period_type": "term",
        "term_code": "second",
        "page_size": "a4_portrait",
        "template": "reports/final_term_2.html",
    },
    "final_annual": {
        "period_type": "annual",
        "page_size": "a4_portrait",
        "template": "reports/final_annual.html",
    },
    "summer_report": {
        "period_type": "summer",
        "page_size": "a4_portrait",
        "template": "reports/summer_report.html",
    },
}
REPORT_CARD_TEMPLATE_KEYS = tuple(REPORT_CARD_LAYOUTS)

ALLOWED_CONTENT_OVERRIDES = {
    "manager_comment": 1000,
    "family_recommendations": 2000,
    "supplemental_text": 2000,
    "display_title": 120,
    "footer_text": 240,
}
EDITABLE_DOCX_NOTICE = (
    "نسخه قابل ویرایش — اعتبار نهایی با PDF آرشیوشده سامانه"
)
REPORT_SNAPSHOT_SCHEMA_VERSION = "report-card-v2"


def report_page_size(presentation):
    """Return a safe CSS @page value for a frozen template presentation."""
    if not isinstance(presentation, dict):
        return ALLOWED_REPORT_PAGE_SIZES["a4_portrait"]
    return ALLOWED_REPORT_PAGE_SIZES.get(
        presentation.get("page_size"), ALLOWED_REPORT_PAGE_SIZES["a4_portrait"]
    )


def _decimal_string(value, decimal_places=2):
    return f"{value:.{decimal_places}f}" if value is not None else None


def _json_value(value):
    if isinstance(value, Decimal):
        return str(value)
    if hasattr(value, "isoformat"):
        return value.isoformat()
    if isinstance(value, uuid.UUID):
        return str(value)
    raise TypeError(f"Unsupported report fingerprint value: {type(value).__name__}")


def source_fingerprint(snapshot):
    """Hash only authoritative report inputs, excluding approval/transport metadata."""
    source = deepcopy(snapshot)
    source.pop("official", None)
    source.pop("source_fingerprint", None)
    source.pop("generated_at", None)
    canonical = json.dumps(
        source,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=_json_value,
    )
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def build_family_safe_context(enrollment):
    """Return presentation facts that are safe for a student/family report."""
    from hamamooz.apps.activities.models import ActivityParticipation
    from hamamooz.apps.attendance.models import AttendanceRecord, AttendanceSession
    from hamamooz.apps.evaluations.models import MonthlyEvaluation
    from hamamooz.apps.recommendations.models import Recommendation

    now = timezone.now()
    attendance_records = AttendanceRecord.objects.filter(
        enrollment=enrollment, session__status=AttendanceSession.Status.FINALIZED
    )
    recommendations = (
        Recommendation.objects.filter(
            enrollment=enrollment,
            audience__in=[Recommendation.Audience.PARENT, Recommendation.Audience.STUDENT],
            status=Recommendation.Status.APPROVED,
            approved_at__isnull=False,
        )
        .filter(Q(expires_at__isnull=True) | Q(expires_at__gt=now))
        .order_by("created_at", "id")
    )
    return {
        "attendance": {
            "finalized_session_count": attendance_records.values("session_id").distinct().count(),
            "unexcused_absence_count": attendance_records.filter(
                status=AttendanceRecord.Status.ABSENT_UNEXCUSED
            ).count(),
        },
        "evaluations": [
            {
                "month_no": item.month_no,
                "framework_version": item.framework_version,
                "metric_scores": {
                    score.metric_code: score.value for score in item.metric_scores.all()
                },
            }
            for item in MonthlyEvaluation.objects.filter(enrollment=enrollment)
            .prefetch_related("metric_scores")
            .order_by("month_no", "framework_version")
        ],
        "activities": [
            {
                "title": item.activity.title,
                "kind": item.activity.kind,
                "status": item.status,
                "result": item.result,
                "placement": item.placement,
            }
            for item in ActivityParticipation.objects.filter(enrollment=enrollment)
            .select_related("activity")
            .order_by("activity__title", "id")
        ],
        "approved_recommendations": [
            {
                "id": str(item.id),
                "audience": item.audience,
                "approved_text": item.approved_text,
            }
            for item in recommendations
        ],
    }


def _score_chart_svg(subjects):
    rows = [row for row in subjects if row.get("average") is not None]
    if not rows:
        return ""
    width = 540
    row_height = 28
    height = 28 + len(rows) * row_height
    bars = []
    for index, row in enumerate(rows):
        try:
            score = max(Decimal("0"), min(Decimal("20"), Decimal(str(row["average"]))))
        except (ArithmeticError, ValueError):
            continue
        bar_width = int(score / Decimal("20") * Decimal("360"))
        y = 18 + index * row_height
        title = escape(str(row.get("title", "")))
        value = escape(str(row["average"]))
        bars.append(
            f'<text x="530" y="{y + 12}" text-anchor="end">{title}</text>'
            f'<rect x="110" y="{y}" width="{bar_width}" height="15" rx="3" />'
            f'<text x="100" y="{y + 12}" text-anchor="end">{value}</text>'
        )
    return (
        f'<svg viewBox="0 0 {width} {height}" role="img" '
        'aria-label="نمودار مقایسه نمرات" xmlns="http://www.w3.org/2000/svg">'
        '<g fill="#b58a32" font-family="Vazirmatn, sans-serif" font-size="11">'
        + "".join(bars)
        + "</g></svg>"
    )


def _evaluation_analysis(context):
    evaluations = context.get("evaluations", [])
    if not evaluations:
        return {
            "evaluation_radar_svg": "",
            "evaluation_framework_version": "",
            "strengths": [],
            "weaknesses": [],
            "evaluation_trend": None,
        }
    latest = evaluations[-1]
    framework_version = latest.get("framework_version", "")
    metric_items = []
    for code, raw_value in sorted(latest.get("metric_scores", {}).items()):
        try:
            value = max(Decimal("0"), min(Decimal("5"), Decimal(str(raw_value))))
        except (ArithmeticError, ValueError):
            continue
        metric_items.append((str(code), value))
    radar_svg = ""
    if len(metric_items) >= 3:
        center = Decimal("110")
        radius = Decimal("82")
        points = []
        axes = []
        for index, (code, value) in enumerate(metric_items):
            angle = -math.pi / 2 + 2 * math.pi * index / len(metric_items)
            axis_x = float(center) + float(radius) * math.cos(angle)
            axis_y = float(center) + float(radius) * math.sin(angle)
            score_radius = float(radius * value / Decimal("5"))
            score_x = float(center) + score_radius * math.cos(angle)
            score_y = float(center) + score_radius * math.sin(angle)
            points.append(f"{score_x:.1f},{score_y:.1f}")
            label_x = float(center) + float(radius + Decimal("16")) * math.cos(angle)
            label_y = float(center) + float(radius + Decimal("16")) * math.sin(angle)
            axes.append(
                f'<line x1="110" y1="110" x2="{axis_x:.1f}" y2="{axis_y:.1f}" />'
                f'<text x="{label_x:.1f}" y="{label_y:.1f}" stroke="none" '
                f'fill="#102f50">{escape(code)}</text>'
            )
        radar_svg = (
            '<svg viewBox="0 0 220 220" role="img" '
            'aria-label="نمودار راداری ارزیابی">'
            '<g stroke="#9fb0bf" fill="none">'
            + "".join(axes)
            + '</g><polygon points="'
            + " ".join(points)
            + '" fill="#b58a3266" stroke="#102f50" stroke-width="2" />'
            '<g fill="#102f50" font-size="8" text-anchor="middle"></g></svg>'
        )
    descending = sorted(metric_items, key=lambda item: (-item[1], item[0]))
    ascending = sorted(metric_items, key=lambda item: (item[1], item[0]))
    same_framework = [
        item for item in evaluations if item.get("framework_version") == framework_version
    ]
    trend = None
    if len(same_framework) >= 2:
        first_values = [
            Decimal(str(value)) for value in same_framework[0]["metric_scores"].values()
        ]
        last_values = [
            Decimal(str(value)) for value in same_framework[-1]["metric_scores"].values()
        ]
        if first_values and last_values:
            delta = sum(last_values, Decimal("0")) / len(last_values) - sum(
                first_values, Decimal("0")
            ) / len(first_values)
            trend = {
                "delta": _decimal_string(delta),
                "label": (
                    "رو به رشد" if delta > 0 else ("کاهشی" if delta < 0 else "بدون تغییر")
                ),
            }
    return {
        "evaluation_radar_svg": radar_svg,
        "evaluation_framework_version": framework_version,
        "strengths": [
            {"metric_code": code, "value": str(value)} for code, value in descending[:3]
        ],
        "weaknesses": [
            {"metric_code": code, "value": str(value)} for code, value in ascending[:3]
        ],
        "evaluation_trend": trend,
    }


def _attach_family_context(report, enrollment):
    context = build_family_safe_context(enrollment)
    report["product_context"] = context
    report.update(_evaluation_analysis(context))


def build_student_snapshot(enrollment, term, *, recalculate=True):
    if recalculate:
        if enrollment.status == Enrollment.Status.ACTIVE:
            recalculate_class_term(enrollment.class_section, term)
        else:
            term_result = calculate_enrollment_term(enrollment, term)
            if term_result.class_rank is not None:
                term_result.class_rank = None
                term_result.save(update_fields=["class_rank", "updated_at"])
    policy = get_policy(enrollment)
    term_result = TermResult.objects.get(enrollment=enrollment, term=term)
    offerings = list(
        CourseOffering.objects.filter(
            class_section=enrollment.class_section, term=term, is_active=True
        ).select_related("grade_subject__subject")
    )
    result_map = {
        row.course_offering_id: row
        for row in SubjectResult.objects.filter(
            enrollment=enrollment, course_offering__in=offerings
        )
    }
    category_buckets = defaultdict(lambda: defaultdict(lambda: [Decimal("0"), Decimal("0")]))
    scores = Score.objects.filter(
        enrollment=enrollment,
        assessment__course_offering__in=offerings,
        assessment__status__in=[Assessment.Status.APPROVED, Assessment.Status.LOCKED],
    ).select_related("assessment", "assessment__assessment_type")
    for score in scores:
        value = normalized_score(score, policy)
        if value is None:
            continue
        offering_id = score.assessment.course_offering_id
        category = score.assessment.assessment_type.category
        category_buckets[offering_id][category][0] += value * score.assessment.weight
        category_buckets[offering_id][category][1] += score.assessment.weight

    subjects = []
    for offering in offerings:
        result = result_map.get(offering.id)
        categories = {
            category: quantize(total / weight, policy) if weight else None
            for category, (total, weight) in category_buckets[offering.id].items()
        }
        subjects.append(
            {
                "title": offering.grade_subject.subject.title,
                "coefficient": str(offering.grade_subject.coefficient),
                "continuous": _decimal_string(categories.get("continuous"), policy.decimal_places),
                "midterm": _decimal_string(categories.get("midterm"), policy.decimal_places),
                "final": _decimal_string(categories.get("final"), policy.decimal_places),
                "average": _decimal_string(
                    result.average if result else None, policy.decimal_places
                ),
                "passed": result.passed if result else False,
            }
        )
    school = enrollment.school
    logo_url = school.logo.url if school.logo else ""
    report = {
        "organization": {
            "name": school.organization.name,
        },
        "school": {
            "name": school.official_name or school.name,
            "branch": school.name if school.official_name else "",
            "address": school.address,
            "phone": school.phone,
            "manager": school.manager_name,
            "logo_url": logo_url,
        },
        "student": {
            "full_name": enrollment.student.full_name,
            "national_id": enrollment.student.national_id,
            "student_number": enrollment.student_number,
            "photo_url": enrollment.student.photo.url if enrollment.student.photo else "",
        },
        "academic": {
            "year": enrollment.academic_year.title,
            "term": term.title,
            "grade": enrollment.grade_level.title,
            "class": enrollment.class_section.title,
        },
        "subjects": subjects,
        "summary": {
            "average": _decimal_string(term_result.average, policy.decimal_places),
            "class_rank": term_result.class_rank,
            "grade_rank": getattr(term_result, "grade_rank", None),
            "school_rank": getattr(term_result, "school_rank", None),
            "class_population": getattr(term_result, "class_population", None),
            "grade_population": getattr(term_result, "grade_population", None),
            "school_population": getattr(term_result, "school_population", None),
            "passed": term_result.passed,
            "status_label": "قبول" if term_result.passed else "نیازمند بررسی",
            "formula_version": term_result.formula_version,
        },
    }
    report["chart_svg"] = _score_chart_svg(subjects)
    return report


def build_report_snapshot(report_type, term, enrollment=None, class_section=None):
    if report_type == ReportArchive.ReportType.STUDENT_REPORT_CARD:
        return {"reports": [build_student_snapshot(enrollment, term)]}
    enrollments = list(
        Enrollment.all_objects.filter(
            class_section=class_section,
            enrolled_on__lte=term.ends_on,
            is_deleted=False,
        )
        .filter(Q(left_on__isnull=True) | Q(left_on__gte=term.starts_on))
        .select_related("student", "school", "academic_year", "grade_level", "class_section")
    )
    recalculate_class_term(class_section, term)
    return {
        "reports": [build_student_snapshot(item, term, recalculate=False) for item in enrollments]
    }


def _rank_visibility(school, academic_year):
    from hamamooz.apps.academics.calculations import get_academic_report_settings

    report_settings = get_academic_report_settings(school, academic_year)
    return {
        "class": report_settings.show_class_rank,
        "grade": report_settings.show_grade_rank,
        "school": report_settings.show_school_rank,
    }, {
        "first_term": str(report_settings.first_term_weight),
        "second_term": str(report_settings.second_term_weight),
        "revision": report_settings.revision,
    }


def _report_scope_enrollments(*, enrollment=None, class_section=None, term=None):
    if enrollment is not None:
        return [enrollment]
    queryset = Enrollment.all_objects.filter(class_section=class_section, is_deleted=False)
    if term is not None:
        queryset = queryset.filter(enrolled_on__lte=term.ends_on).filter(
            Q(left_on__isnull=True) | Q(left_on__gte=term.starts_on)
        )
    else:
        queryset = queryset.filter(
            academic_year=class_section.academic_year,
            status=Enrollment.Status.ACTIVE,
        )
    return list(
        queryset.select_related(
            "student", "school__organization", "academic_year", "grade_level", "class_section"
        ).order_by("student__last_name", "student__first_name", "id")
    )


def _build_term_card_reports(*, term, enrollment=None, class_section=None):
    scope_class = enrollment.class_section if enrollment else class_section
    try:
        from hamamooz.apps.academics.calculations import recalculate_school_term
    except ImportError:  # Backward-compatible during additive deployment.
        recalculate_class_term(scope_class, term)
    else:
        recalculate_school_term(scope_class.school, term)
    enrollments = _report_scope_enrollments(
        enrollment=enrollment, class_section=class_section, term=term
    )
    visibility, weights = _rank_visibility(scope_class.school, scope_class.academic_year)
    reports = []
    for item in enrollments:
        report = build_student_snapshot(item, term, recalculate=False)
        report["rank_visibility"] = visibility
        report["term_weights"] = weights
        _attach_family_context(report, item)
        reports.append(report)
    return reports


def _annual_term_result_map(annual_result):
    from hamamooz.apps.organizations.models import Term

    anchor = annual_result.enrollment
    historical = list(
        Enrollment.all_objects.filter(
            student_id=anchor.student_id,
            school_id=anchor.school_id,
            academic_year_id=anchor.academic_year_id,
            grade_level_id=anchor.grade_level_id,
            is_deleted=False,
        ).select_related("class_section")
    )
    terms = list(
        Term.objects.filter(
            academic_year_id=anchor.academic_year_id,
            code__in=[Term.Code.FIRST, Term.Code.SECOND],
        ).order_by("order", "pk")
    )
    eligible_pairs = {
        (segment.class_section_id, term.id)
        for segment in historical
        for term in terms
        if segment.enrolled_on <= term.ends_on
        and (segment.left_on is None or segment.left_on >= term.starts_on)
    }
    rows = SubjectResult.objects.filter(
        enrollment_id__in=[segment.id for segment in historical],
        course_offering__term_id__in=[term.id for term in terms],
    ).select_related(
        "course_offering__term", "course_offering__grade_subject__subject", "enrollment"
    )
    selected = {}
    for row in rows.order_by(
        "course_offering__term__order", "enrollment__enrolled_on", "calculated_at", "pk"
    ):
        if (row.enrollment.class_section_id, row.course_offering.term_id) not in eligible_pairs:
            continue
        selected[(row.course_offering.grade_subject_id, row.course_offering.term.code)] = row
    return selected


def _build_annual_student_snapshot(enrollment):
    from hamamooz.apps.academics.calculations import (
        calculate_enrollment_annual,
        get_academic_report_settings,
        get_policy,
    )

    result = calculate_enrollment_annual(enrollment)
    enrollment = result.enrollment
    policy = get_policy(enrollment)
    report_settings = get_academic_report_settings(enrollment.school, enrollment.academic_year)
    by_subject_term = _annual_term_result_map(result)
    subject_results = result.subject_results.select_related(
        "grade_subject__subject"
    ).order_by("grade_subject__subject__title")
    subjects = []
    for row in subject_results:
        first = by_subject_term.get((row.grade_subject_id, "first"))
        second = by_subject_term.get((row.grade_subject_id, "second"))
        subjects.append({
            "title": row.grade_subject.subject.title,
            "coefficient": str(row.grade_subject.coefficient),
            "continuous": None,
            "midterm": None,
            "final": None,
            "first_term_score": _decimal_string(
                first.average if first else None, policy.decimal_places
            ),
            "second_term_score": _decimal_string(
                second.average if second else None, policy.decimal_places
            ),
            "average": _decimal_string(row.average, policy.decimal_places),
            "complete": row.complete,
            "passed": row.passed,
        })
    school = enrollment.school
    report = {
        "organization": {"name": school.organization.name},
        "school": {
            "name": school.official_name or school.name,
            "branch": school.name if school.official_name else "",
            "address": school.address,
            "phone": school.phone,
            "manager": school.manager_name,
            "logo_url": school.logo.url if school.logo else "",
        },
        "student": {
            "full_name": enrollment.student.full_name,
            "national_id": enrollment.student.national_id,
            "student_number": enrollment.student_number,
            "photo_url": enrollment.student.photo.url if enrollment.student.photo else "",
        },
        "academic": {
            "year": enrollment.academic_year.title,
            "term": "سالانه",
            "grade": enrollment.grade_level.title,
            "class": enrollment.class_section.title,
        },
        "subjects": subjects,
        "summary": {
            "average": _decimal_string(result.average, policy.decimal_places),
            "class_rank": result.class_rank,
            "grade_rank": result.grade_rank,
            "school_rank": result.school_rank,
            "class_population": result.class_population,
            "grade_population": result.grade_population,
            "school_population": result.school_population,
            "complete": result.complete,
            "passed": result.passed,
            "status_label": "قبول" if result.passed else "نیازمند بررسی",
            "formula_version": result.formula_version,
        },
        "rank_visibility": {
            "class": report_settings.show_class_rank,
            "grade": report_settings.show_grade_rank,
            "school": report_settings.show_school_rank,
        },
        "term_weights": {
            "first_term": str(report_settings.first_term_weight),
            "second_term": str(report_settings.second_term_weight),
            "revision": report_settings.revision,
        },
    }
    _attach_family_context(report, enrollment)
    report["chart_svg"] = _score_chart_svg(subjects)
    return report


def _build_annual_card_reports(*, enrollment=None, class_section=None):
    from hamamooz.apps.academics.calculations import recalculate_school_annual

    scope_class = enrollment.class_section if enrollment else class_section
    recalculate_school_annual(scope_class.school, scope_class.academic_year)
    return [
        _build_annual_student_snapshot(item)
        for item in _report_scope_enrollments(enrollment=enrollment, class_section=class_section)
    ]


def _validate_annual_locked_sources(*, enrollment=None, class_section=None):
    from rest_framework.exceptions import ValidationError as DRFValidationError

    from hamamooz.apps.academics.calculations import calculate_enrollment_annual
    from hamamooz.apps.academics.services import validate_score_completeness

    invalid = set()
    for item in _report_scope_enrollments(enrollment=enrollment, class_section=class_section):
        annual_result = calculate_enrollment_annual(item)
        if not annual_result.complete:
            invalid.add(item.student.full_name)
            continue
        selected = _annual_term_result_map(annual_result)
        required_subject_ids = set(
            annual_result.subject_results.values_list("grade_subject_id", flat=True)
        )
        if any(
            (subject_id, term_code) not in selected
            for subject_id in required_subject_ids
            for term_code in ("first", "second")
        ):
            invalid.add(item.student.full_name)
            continue
        for row in selected.values():
            if row.course_offering.grade_subject_id not in required_subject_ids:
                continue
            assessments = list(row.course_offering.assessments.all())
            if not assessments or any(
                assessment.status != Assessment.Status.LOCKED for assessment in assessments
            ):
                invalid.add(row.course_offering.grade_subject.subject.title)
                continue
            try:
                for assessment in assessments:
                    validate_score_completeness(assessment)
            except DRFValidationError:
                invalid.add(row.course_offering.grade_subject.subject.title)
    if invalid:
        raise ValueError(
            "برای صدور سالانه، ارزیابی‌های هر دو نوبت باید کامل و قفل باشند: "
            + "، ".join(sorted(invalid)[:10])
        )


def _summer_preview_result(registration, exam=None):
    from hamamooz.apps.summers.models import SummerSubjectScore

    resolved_exam = exam or registration.program.exams.order_by("-exam_date", "id").first()
    course_registrations = list(
        registration.course_registrations.select_related("course__subject").order_by(
            "course__subject__title"
        )
    )
    scores = {}
    if resolved_exam:
        scores = {
            item.course_registration_id: item
            for item in SummerSubjectScore.objects.filter(
                exam=resolved_exam, course_registration__in=course_registrations
            )
        }
    courses = []
    for item in course_registrations:
        score = scores.get(item.id)
        courses.append(
            {
                "course_registration_id": str(item.id),
                "subject_id": str(item.course.subject_id),
                "subject_title": item.course.subject.title,
                "coefficient": item.course.subject.default_coefficient,
                "score": score.value if score else None,
            }
        )
    return {
        "program_id": str(registration.program_id),
        "registration_id": str(registration.id),
        "enrollment_id": str(registration.enrollment_id),
        "exam_id": str(resolved_exam.id) if resolved_exam else None,
        "courses": courses,
        "average": None,
        "pass_threshold": registration.program.pass_threshold,
        "passed": None,
    }


def _build_summer_student_snapshot(registration, exam=None, *, require_ready=False):
    from rest_framework.exceptions import ValidationError as DRFValidationError

    from hamamooz.apps.summers.services import summer_registration_result

    warnings = []
    try:
        result = summer_registration_result(registration, exam)
    except DRFValidationError as exc:
        if require_ready:
            raise
        result = _summer_preview_result(registration, exam)
        warnings.append(str(exc.detail))
    enrollment = registration.enrollment
    threshold = result["pass_threshold"]
    subjects = [
        {
            "title": item["subject_title"],
            "coefficient": str(item["coefficient"]),
            "continuous": None,
            "midterm": None,
            "final": _decimal_string(item["score"]),
            "average": _decimal_string(item["score"]),
            "passed": (
                None
                if threshold is None or item["score"] is None
                else item["score"] >= threshold
            ),
        }
        for item in result["courses"]
    ]
    school = enrollment.school
    report = {
        "organization": {"name": school.organization.name},
        "school": {
            "name": school.official_name or school.name,
            "branch": school.name if school.official_name else "",
            "address": school.address,
            "phone": school.phone,
            "manager": school.manager_name,
            "logo_url": school.logo.url if school.logo else "",
        },
        "student": {
            "full_name": enrollment.student.full_name,
            "national_id": enrollment.student.national_id,
            "student_number": enrollment.student_number,
            "photo_url": enrollment.student.photo.url if enrollment.student.photo else "",
        },
        "academic": {
            "year": enrollment.academic_year.title,
            "term": registration.program.title,
            "grade": enrollment.grade_level.title,
            "class": enrollment.class_section.title,
        },
        "subjects": subjects,
        "summary": {
            "average": _decimal_string(result["average"]),
            "passed": result["passed"],
            "status_label": (
                ""
                if result["passed"] is None
                else ("قبول" if result["passed"] else "مردود")
            ),
            "formula_version": "summer-direct-score-v1",
        },
        "pass_threshold": _decimal_string(threshold),
        "show_status": threshold is not None,
        "rank_visibility": {"class": False, "grade": False, "school": False},
        "warnings": warnings,
    }
    return report, result.get("exam_id"), warnings


def build_report_card_snapshot(
    template,
    *,
    term=None,
    enrollment=None,
    class_section=None,
    summer_registration=None,
    summer_exam=None,
    require_ready=False,
):
    """Build one of the seven safe report families from authoritative domain services."""
    layout_key = template.layout_key
    if layout_key not in REPORT_CARD_LAYOUTS:
        raise ValueError("Unknown report-card layout.")
    config = REPORT_CARD_LAYOUTS[layout_key]
    period_type = config["period_type"]
    warnings = []
    if period_type == ReportPeriodType.TERM:
        if term is None or term.code != config["term_code"]:
            raise ValueError("The selected term does not match the report layout.")
        if require_ready:
            validate_report_card_readiness(
                template.report_type,
                term=term,
                enrollment=enrollment,
                class_section=class_section,
            )
        reports = _build_term_card_reports(
            term=term, enrollment=enrollment, class_section=class_section
        )
        if any(
            not report["subjects"]
            or any(subject.get("average") is None for subject in report["subjects"])
            for report in reports
        ):
            warnings.append("نتایج نوبت برای همه دروس کامل نیست.")
        period = {"type": "term", "id": str(term.id), "label": term.title}
    elif period_type == ReportPeriodType.ANNUAL:
        reports = _build_annual_card_reports(enrollment=enrollment, class_section=class_section)
        incomplete = [item for item in reports if not item["summary"].get("complete")]
        if incomplete:
            warnings.append(
                "نتایج سالانه برای همه دروس و هر دو نوبت کامل نیست."
            )
            if require_ready:
                raise ValueError(warnings[0])
        if require_ready:
            _validate_annual_locked_sources(enrollment=enrollment, class_section=class_section)
        academic_year = (enrollment or class_section).academic_year
        period = {
            "type": "annual",
            "id": str(academic_year.id),
            "label": academic_year.title,
        }
    else:
        if summer_registration is None:
            raise ValueError("Summer registration is required.")
        report, exam_id, warnings = _build_summer_student_snapshot(
            summer_registration, summer_exam, require_ready=require_ready
        )
        reports = [report]
        period = {
            "type": "summer",
            "id": str(summer_registration.program_id),
            "exam_id": exam_id,
            "label": summer_registration.program.title,
        }
    scope_enrollment = (
        summer_registration.enrollment if summer_registration is not None else enrollment
    )
    scope_object = scope_enrollment or class_section
    school = scope_object.school
    academic_year = scope_object.academic_year
    snapshot = {
        "schema_version": REPORT_SNAPSHOT_SCHEMA_VERSION,
        "layout_key": layout_key,
        "template_key": layout_key,
        "report_type": template.report_type,
        "report_family": (
            "summer"
            if layout_key == ReportLayoutKey.SUMMER_REPORT
            else ("analytical" if layout_key.startswith("analytical_") else "final")
        ),
        "organization_id": str(school.organization_id),
        "school_id": str(school.id),
        "academic_year_id": str(academic_year.id),
        "student_id": (
            str(scope_enrollment.student_id) if scope_enrollment is not None else None
        ),
        "enrollment_id": str(scope_enrollment.id) if scope_enrollment is not None else None,
        "class_section_id": (
            str(scope_enrollment.class_section_id)
            if scope_enrollment is not None
            else str(class_section.id)
        ),
        "summer_program_id": (
            str(summer_registration.program_id) if summer_registration is not None else None
        ),
        "summer_registration_id": (
            str(summer_registration.id) if summer_registration is not None else None
        ),
        "period_type": period["type"],
        "period_label": period["label"],
        "period": period,
        "reports": reports,
        "warnings": warnings,
        "template": {
            "id": str(template.id),
            "code": template.code,
            "title": template.title,
            "blocks": list(template.blocks),
            "presentation": dict(template.presentation),
            "output_format": template.output_format,
        },
    }
    snapshot["source_fingerprint"] = source_fingerprint(snapshot)
    return snapshot


def validate_report_card_readiness(
    report_type,
    *,
    term=None,
    enrollment=None,
    class_section=None,
    summer_registration=None,
    summer_exam=None,
):
    """Revalidate locked official facts without trusting a stored snapshot."""
    from rest_framework.exceptions import ValidationError as DRFValidationError

    from hamamooz.apps.academics.services import validate_score_completeness

    if summer_registration is not None:
        from hamamooz.apps.summers.services import validate_summer_report_readiness

        return validate_summer_report_readiness(summer_registration, summer_exam)
    if term is None:
        raise ValueError("A term is required for term readiness validation.")
    scope_class = class_section or enrollment.class_section
    offerings = CourseOffering.objects.filter(
        class_section=scope_class, term=term, is_active=True
    ).select_related("grade_subject__subject")
    if not offerings.exists():
        raise ValueError(
            "برای این کلاس و نوبت هیچ درس فعالی تعریف نشده است."
        )
    should_validate_roster = (
        report_type == ReportArchive.ReportType.CLASS_REPORT_CARDS
        or enrollment.status == Enrollment.Status.ACTIVE
    )
    incomplete = []
    for offering in offerings:
        assessments = list(offering.assessments.all())
        is_incomplete = not assessments or any(
            assessment.status != Assessment.Status.LOCKED for assessment in assessments
        )
        if not is_incomplete and should_validate_roster:
            try:
                for assessment in assessments:
                    validate_score_completeness(assessment)
            except DRFValidationError:
                is_incomplete = True
        if is_incomplete:
            incomplete.append(offering.grade_subject.subject.title)
    if incomplete:
        titles = "، ".join(incomplete[:10])
        raise ValueError(
            f"برای صدور رسمی، همه ارزیابی‌های دروس باید قفل شوند: {titles}"
        )
    return True


def render_report_html(snapshot, *, preview=False):
    template = snapshot.get("template", {})
    layout_key = snapshot.get("layout_key")
    if layout_key in REPORT_CARD_LAYOUTS:
        template_name = REPORT_CARD_LAYOUTS[layout_key]["template"]
        page_size = ALLOWED_REPORT_PAGE_SIZES[REPORT_CARD_LAYOUTS[layout_key]["page_size"]]
    else:
        template_name = "reports/report_card.html"
        page_size = report_page_size(template.get("presentation"))
    return render_to_string(
        template_name,
        {
            "reports": snapshot["reports"],
            "preview": preview,
            "blocks": template.get("blocks", ALLOWED_REPORT_BLOCKS),
            "overrides": snapshot.get("content_overrides", {}),
            "page_size": page_size,
            "layout_key": layout_key,
            "template_title": template.get("title", ""),
            "period": snapshot.get("period", {}),
            "official": snapshot.get("official", {}),
            "warnings": snapshot.get("warnings", []),
            "generated_at": timezone.now(),
        },
    )


def _local_media_file_url(url):
    if not url or not url.startswith(settings.MEDIA_URL):
        return url
    relative = url.removeprefix(settings.MEDIA_URL).lstrip("/")
    media_root = Path(settings.MEDIA_ROOT).resolve()
    candidate = (media_root / relative).resolve()
    if not candidate.is_relative_to(media_root):
        return ""
    return candidate.as_uri()


def _pdf_snapshot(snapshot):
    rendered = deepcopy(snapshot)
    for report in rendered.get("reports", []):
        report["school"]["logo_url"] = _local_media_file_url(report["school"].get("logo_url", ""))
        report["student"]["photo_url"] = _local_media_file_url(
            report["student"].get("photo_url", "")
        )
    return rendered


def render_report_pdf(snapshot):
    from weasyprint import HTML

    html = render_report_html(_pdf_snapshot(snapshot))
    return HTML(
        string=html,
        base_url=Path(settings.BASE_DIR).as_uri(),
    ).write_pdf(presentational_hints=False)


def render_report_docx(snapshot):
    """Generate a genuinely editable RTL Word document from the approved snapshot."""
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    def rtl(paragraph, *, bold=False, size=10):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        properties = paragraph._p.get_or_add_pPr()
        if properties.find(qn("w:bidi")) is None:
            properties.append(OxmlElement("w:bidi"))
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "Vazirmatn"
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn("w:cs"), "Vazirmatn")

    document = Document()
    section = document.sections[0]
    if snapshot.get("layout_key", "").startswith("analytical_"):
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Cm(42), Cm(29.7)
    else:
        section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.right_margin = section.left_margin = Cm(1.5)
    notice = document.add_paragraph(EDITABLE_DOCX_NOTICE)
    rtl(notice, bold=True, size=9)
    official = snapshot.get("official", {})
    metadata = document.add_paragraph(
        f"کد رهگیری: {official.get('tracking_code', '-')} | "
        f"نسخه: {official.get('version', '-')} | "
        f"تاریخ صدور: {official.get('generated_at') or official.get('approved_at') or '-'}"
    )
    rtl(metadata, size=9)
    for index, report in enumerate(snapshot.get("reports", [])):
        if index:
            document.add_page_break()
        is_annual = snapshot.get("period", {}).get("type") == ReportPeriodType.ANNUAL
        is_summer = snapshot.get("period", {}).get("type") == "summer"
        default_title = "کارنامه دوره تابستان" if is_summer else "کارنامه تحصیلی"
        title = snapshot.get("content_overrides", {}).get(
            "display_title",
            snapshot.get("template", {}).get("title", default_title),
        )
        heading = document.add_heading(title, level=1)
        rtl(heading, bold=True, size=16)
        identity = document.add_paragraph(
            f"نام دانش‌آموز: {report['student']['full_name']} | "
            f"کد ملی: {report['student']['national_id']} | "
            f"شماره دانش‌آموزی: {report['student']['student_number']}"
        )
        rtl(identity, size=10)
        academic = document.add_paragraph(
            f"مدرسه: {report['school']['name']} | "
            f"سال تحصیلی: {report['academic']['year']} | "
            f"پایه: {report['academic']['grade']} | کلاس: {report['academic']['class']} | "
            f"دوره: {snapshot.get('period', {}).get('label', report['academic']['term'])}"
        )
        rtl(academic, size=10)
        if is_summer:
            summer_exam = document.add_paragraph("نتایج آزمون جامع تابستان")
            rtl(summer_exam, bold=True, size=11)
        show_status = report.get("show_status", True)
        columns = ["ردیف", "درس"]
        if not is_summer:
            columns.append("ضریب")
        if is_annual:
            columns.extend(["نوبت اول", "نوبت دوم", "میانگین وزنی"])
        else:
            columns.append("نمره آزمون جامع" if is_summer else "نمره / میانگین")
        if show_status:
            columns.append("وضعیت")
        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, columns, strict=True):
            cell.text = label
            rtl(cell.paragraphs[0], bold=True, size=9)
        for row_number, subject in enumerate(report.get("subjects", []), start=1):
            cells = table.add_row().cells
            values = [
                str(row_number),
                str(subject.get("title", "")),
            ]
            if not is_summer:
                values.append(str(subject.get("coefficient", "")))
            if is_annual:
                values.extend(
                    [
                        str(subject.get("first_term_score") or "-"),
                        str(subject.get("second_term_score") or "-"),
                        str(subject.get("average") or "-"),
                    ]
                )
            else:
                values.append(str(subject.get("average") or "-"))
            if show_status:
                passed = subject.get("passed")
                values.append(
                    "-" if passed is None else ("قبول" if passed else "نیازمند بررسی")
                )
            for cell, value in zip(cells, values, strict=True):
                cell.text = value
                rtl(cell.paragraphs[0], size=9)
        summary_parts = [f"معدل: {report.get('summary', {}).get('average') or '-'}"]
        if show_status:
            summary_parts.append(
                f"نتیجه: {report.get('summary', {}).get('status_label') or '-'}"
            )
        summary = document.add_paragraph(" | ".join(summary_parts))
        rtl(summary, bold=True, size=10)
        rank_labels = {
            "class": "رتبه کلاس",
            "grade": "رتبه پایه",
            "school": "رتبه مدرسه",
        }
        rank_parts = []
        for scope, label in rank_labels.items():
            if report.get("rank_visibility", {}).get(scope):
                rank = report.get("summary", {}).get(f"{scope}_rank") or "-"
                population = report.get("summary", {}).get(f"{scope}_population") or "-"
                rank_parts.append(f"{label}: {rank} از {population}")
        if rank_parts:
            ranks = document.add_paragraph(" | ".join(rank_parts))
            rtl(ranks, size=10)
        blocks = set(snapshot.get("template", {}).get("blocks", ALLOWED_REPORT_BLOCKS))
        product_context = report.get("product_context", {})
        if "attendance_summary" in blocks:
            attendance = product_context.get("attendance", {})
            attendance_paragraph = document.add_paragraph(
                "جلسات نهایی حضور: "
                f"{attendance.get('finalized_session_count', 0)} | "
                "غیبت غیرموجه: "
                f"{attendance.get('unexcused_absence_count', 0)}"
            )
            rtl(attendance_paragraph, size=10)
        if snapshot.get("layout_key", "").startswith("analytical_"):
            strengths = "، ".join(
                f"{item['metric_code']} ({item['value']})"
                for item in report.get("strengths", [])
            )
            weaknesses = "، ".join(
                f"{item['metric_code']} ({item['value']})"
                for item in report.get("weaknesses", [])
            )
            analysis = document.add_paragraph(
                f"نقاط قوت: {strengths or '-'} | زمینه‌های قابل بهبود: {weaknesses or '-'}"
            )
            rtl(analysis, size=10)
        overrides = snapshot.get("content_overrides", {})
        for key in ("manager_comment", "family_recommendations", "supplemental_text"):
            if overrides.get(key):
                paragraph = document.add_paragraph(str(overrides[key]))
                rtl(paragraph, size=10)
        if "recommendations" in blocks and not overrides.get("family_recommendations"):
            for recommendation in product_context.get("approved_recommendations", []):
                paragraph = document.add_paragraph(recommendation["approved_text"])
                rtl(paragraph, size=10)
        signatures = document.add_paragraph(
            "امضای معاون آموزشی          "
            f"امضا و مهر مدیر: {report.get('school', {}).get('manager') or '-'}"
        )
        rtl(signatures, bold=True, size=10)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _report_extended_context(enrollment):
    """Backward-compatible alias for the tightened family-safe context."""
    return build_family_safe_context(enrollment)


def build_draft_snapshot(
    template,
    *,
    term=None,
    enrollment=None,
    class_section=None,
    summer_registration=None,
    summer_exam=None,
):
    """Freeze report inputs at draft creation; official readiness is checked later."""
    if template.layout_key:
        return build_report_card_snapshot(
            template,
            term=term,
            enrollment=enrollment,
            class_section=class_section,
            summer_registration=summer_registration,
            summer_exam=summer_exam,
        )
    snapshot = build_report_snapshot(
        template.report_type,
        term,
        enrollment=enrollment,
        class_section=class_section,
    )
    enrollments = (
        [enrollment]
        if enrollment
        else list(
            Enrollment.all_objects.filter(
                class_section=class_section,
                enrolled_on__lte=term.ends_on,
                is_deleted=False,
            )
            .filter(Q(left_on__isnull=True) | Q(left_on__gte=term.starts_on))
            .select_related("student")
        )
    )
    for report, subject in zip(snapshot["reports"], enrollments, strict=True):
        report["product_context"] = _report_extended_context(subject)
    snapshot["template"] = {
        "id": str(template.id),
        "code": template.code,
        "blocks": list(template.blocks),
        "presentation": template.presentation,
        "output_format": template.output_format,
    }
    snapshot["source_fingerprint"] = source_fingerprint(snapshot)
    return snapshot


def _draft_source_kwargs(draft):
    return {
        "term": draft.term,
        "enrollment": draft.enrollment,
        "class_section": draft.class_section,
        "summer_registration": draft.summer_registration,
        "summer_exam": draft.summer_exam,
    }


def revalidate_report_draft(draft, *, require_ready=True):
    """Reject a draft if any authoritative source changed after its snapshot."""
    if not draft.layout_key:
        if require_ready:
            validate_report_card_readiness(
                draft.template.report_type,
                term=draft.term,
                enrollment=draft.enrollment,
                class_section=draft.class_section,
            )
        return draft.snapshot
    if draft.template.layout_key != draft.layout_key:
        raise ValueError("The report template changed after draft creation.")
    current = build_report_card_snapshot(
        draft.template,
        require_ready=require_ready,
        **_draft_source_kwargs(draft),
    )
    expected = draft.source_fingerprint or draft.snapshot.get("source_fingerprint", "")
    if not expected or current["source_fingerprint"] != expected:
        raise ValueError(
            "Report sources changed after preview; create a fresh draft before approval."
        )
    return current


def _version_scope(queryset, draft):
    queryset = queryset.filter(
        school=draft.school,
        academic_year=draft.academic_year,
        layout_key=draft.layout_key,
    )
    if draft.summer_registration_id:
        return queryset.filter(summer_registration_id=draft.summer_registration_id)
    if draft.enrollment_id:
        queryset = queryset.filter(enrollment_id=draft.enrollment_id)
    else:
        queryset = queryset.filter(class_section_id=draft.class_section_id)
    return queryset.filter(term_id=draft.term_id)


@transaction.atomic
def approve_report_draft(draft_id, *, actor):
    """Freeze approval metadata and reserve the next monotonically increasing version."""
    from hamamooz.apps.organizations.models import School

    draft = (
        ReportDraft.objects.select_for_update()
        .select_related(
            "template",
            "organization",
            "school",
            "academic_year",
            "term",
            "enrollment__student",
            "class_section",
            "summer_program",
            "summer_registration__enrollment__student",
            "summer_exam",
        )
        .get(pk=draft_id)
    )
    if draft.status != ReportDraft.Status.SUBMITTED:
        raise ValueError("Only a submitted report may be approved.")
    # A stable tenant parent exists before the first version and serializes the
    # otherwise-empty MAX(version) scope under concurrent approvals.
    School.objects.select_for_update().get(pk=draft.school_id)
    revalidate_report_draft(draft, require_ready=True)
    draft_versions = _version_scope(ReportDraft.objects.select_for_update(), draft).aggregate(
        value=Max("report_version")
    )["value"]
    archive_versions = _version_scope(ReportArchive.objects.select_for_update(), draft).aggregate(
        value=Max("report_version")
    )["value"]
    version = max(draft_versions or 0, archive_versions or 0) + 1
    tracking_code = f"HMZ-{draft.academic_year.code}-{uuid.uuid4().hex[:12].upper()}"
    approved_at = timezone.now()
    approved_snapshot = deepcopy(draft.snapshot)
    approved_snapshot["content_overrides"] = dict(draft.content_overrides)
    approved_snapshot["official"] = {
        "tracking_code": tracking_code,
        "version": version,
        "approved_at": approved_at.isoformat(),
        "generated_at": approved_at.isoformat(),
        "approved_by": actor.get_full_name() or actor.get_username(),
        "approved_by_id": str(actor.id),
    }
    draft.snapshot = approved_snapshot
    draft.tracking_code = tracking_code
    draft.report_version = version
    draft.status = ReportDraft.Status.APPROVED
    draft.reviewed_by = actor
    draft.reviewed_at = approved_at
    draft.rejection_reason = ""
    draft.full_clean(exclude=["id"])
    draft.save(
        update_fields=[
            "snapshot",
            "tracking_code",
            "report_version",
            "status",
            "reviewed_by",
            "reviewed_at",
            "rejection_reason",
            "updated_at",
        ]
    )
    return draft


def render_report_draft(draft_id):
    """Render exactly the frozen approved snapshot into the immutable archive."""
    with transaction.atomic():
        draft = (
            ReportDraft.objects.select_for_update(of=("self",))
            .select_related(
                "template",
                "organization",
                "school",
                "academic_year",
                "term",
                "enrollment",
                "class_section",
                "summer_program",
                "summer_registration__enrollment",
                "summer_exam",
                "archive",
            )
            .get(pk=draft_id)
        )
        if draft.status == ReportDraft.Status.RENDERED:
            return draft
        if draft.status != ReportDraft.Status.APPROVED:
            raise ValueError("Only an approved report draft may be rendered.")
        if draft.layout_key:
            revalidate_report_draft(draft, require_ready=True)
        render_snapshot = deepcopy(draft.snapshot)
        if not draft.layout_key:
            render_snapshot["content_overrides"] = dict(draft.content_overrides)
        is_new_family = bool(draft.layout_key)
        if draft.archive_id:
            archive = draft.archive
            if archive.status == ReportArchive.Status.COMPLETED:
                draft.status = ReportDraft.Status.RENDERED
                draft.save(update_fields=["status", "updated_at"])
                return draft
            if archive.status == ReportArchive.Status.PROCESSING:
                return draft
            archive.status = ReportArchive.Status.PROCESSING
            archive.started_at = timezone.now()
            archive.completed_at = None
            archive.error_message = ""
            archive.save(
                update_fields=[
                    "status",
                    "started_at",
                    "completed_at",
                    "error_message",
                    "updated_at",
                ]
            )
        else:
            archive = ReportArchive.objects.create(
                organization=draft.organization,
                school=draft.school,
                academic_year=draft.academic_year,
                term=draft.term,
                summer_program=draft.summer_program,
                summer_registration=draft.summer_registration,
                summer_exam=draft.summer_exam,
                report_type=draft.template.report_type,
                layout_key=draft.layout_key,
                status=ReportArchive.Status.PROCESSING,
                enrollment=draft.enrollment,
                class_section=draft.class_section,
                requested_by=draft.created_by,
                output_format=(
                    ReportArchive.OutputFormat.PDF
                    if is_new_family
                    else draft.template.output_format
                ),
                snapshot=render_snapshot,
                source_fingerprint=draft.source_fingerprint,
                tracking_code=draft.tracking_code,
                report_version=draft.report_version,
                approved_by=draft.reviewed_by,
                approved_at=draft.reviewed_at,
                formula_version=(render_snapshot.get("reports") or [{}])[0]
                .get("summary", {})
                .get("formula_version", ""),
                started_at=timezone.now(),
            )
            draft.archive = archive
            draft.save(update_fields=["archive", "updated_at"])
    try:
        editable_name = ""
        if is_new_family:
            output = render_report_pdf(render_snapshot)
            extension = "pdf"
            editable = render_report_docx(render_snapshot)
            editable_filename = f"draft_{draft.id}_{draft.created_at:%Y-%m-%d}.docx"
            archive.editable_output_file.save(
                editable_filename, ContentFile(editable), save=False
            )
            editable_name = archive.editable_output_file.name
        elif draft.template.output_format == draft.template.OutputFormat.DOCX:
            output = render_report_docx(render_snapshot)
            extension = "docx"
        else:
            output = render_report_pdf(render_snapshot)
            extension = "pdf"
        filename = f"draft_{draft.id}_{draft.created_at:%Y-%m-%d}.{extension}"
        archive.output_file.save(filename, ContentFile(output), save=False)
        output_name = archive.output_file.name
        with transaction.atomic():
            archive = ReportArchive.objects.select_for_update().get(pk=archive.pk)
            archive.output_file.name = output_name
            if editable_name:
                archive.editable_output_file.name = editable_name
            archive.status = ReportArchive.Status.COMPLETED
            archive.completed_at = timezone.now()
            archive.error_message = ""
            archive.save(
                update_fields=[
                    "output_file",
                    "editable_output_file",
                    "status",
                    "completed_at",
                    "error_message",
                    "updated_at",
                ]
            )
            draft = ReportDraft.objects.select_for_update().get(pk=draft_id)
            draft.status = ReportDraft.Status.RENDERED
            draft.archive = archive
            draft.save(update_fields=["status", "archive", "updated_at"])
            return draft
    except Exception as exc:
        archive.status = ReportArchive.Status.FAILED
        archive.error_message = str(exc)[:2000]
        archive.completed_at = timezone.now()
        archive.save(update_fields=["status", "error_message", "completed_at", "updated_at"])
        raise


def generate_report(report_id):
    processing_timeout = timedelta(
        minutes=getattr(settings, "REPORT_PROCESSING_TIMEOUT_MINUTES", 30)
    )
    with transaction.atomic():
        report = (
            ReportArchive.objects.select_for_update(of=("self",))
            .select_related("term", "enrollment", "class_section")
            .get(pk=report_id)
        )
        if report.status == ReportArchive.Status.COMPLETED:
            return report
        if report.layout_key:
            raise ValueError(
                "Official report-card families must be rendered from an approved draft."
            )
        if (
            report.status == ReportArchive.Status.PROCESSING
            and report.started_at
            and report.started_at >= timezone.now() - processing_timeout
        ):
            return report
        report.status = ReportArchive.Status.PROCESSING
        report.started_at = timezone.now()
        report.completed_at = None
        report.error_message = ""
        report.save(
            update_fields=[
                "status",
                "started_at",
                "completed_at",
                "error_message",
                "updated_at",
            ]
        )

    stored_name = ""
    try:
        snapshot = build_report_snapshot(
            report.report_type,
            report.term,
            enrollment=report.enrollment,
            class_section=report.class_section,
        )
        pdf = render_report_pdf(snapshot)
        first = snapshot["reports"][0] if snapshot["reports"] else None
        formula_version = first["summary"]["formula_version"] if first else ""
        filename = (
            f"{report.report_type}_"
            f"{report.organization.code}_"
            f"{report.school.code}_"
            f"{report.created_at:%Y-%m-%d}.pdf"
        )
        report.output_file.save(filename, ContentFile(pdf), save=False)
        stored_name = report.output_file.name
        with transaction.atomic():
            locked = ReportArchive.objects.select_for_update().get(pk=report_id)
            if locked.status == ReportArchive.Status.COMPLETED:
                if stored_name and stored_name != locked.output_file.name:
                    report.output_file.storage.delete(stored_name)
                return locked
            locked.output_file = report.output_file
            locked.snapshot = snapshot
            locked.formula_version = formula_version
            locked.status = ReportArchive.Status.COMPLETED
            locked.completed_at = timezone.now()
            locked.error_message = ""
            locked.save()
            return locked
    except Exception as exc:
        if stored_name:
            report.output_file.storage.delete(stored_name)
        with transaction.atomic():
            locked = ReportArchive.objects.select_for_update().get(pk=report_id)
            locked.status = ReportArchive.Status.FAILED
            locked.error_message = str(exc)[:2000]
            locked.completed_at = timezone.now()
            locked.save(update_fields=["status", "error_message", "completed_at", "updated_at"])
        raise
